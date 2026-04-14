#!/usr/bin/env python3
"""
Builds CITL_AI_Workforce_Grant_Document.docx from the Markdown source.
Professional serif layout: Georgia body, headings in deep navy, tables with shading.
"""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1C, 0x34, 0x54)   # deep navy — main headings
MAROON      = RGBColor(0x7B, 0x0C, 0x1A)   # RTC maroon — part headings / accent
GOLD        = RGBColor(0xB8, 0x8A, 0x00)   # gold — app number badges
BODY_TEXT   = RGBColor(0x1E, 0x1E, 0x1E)   # near-black body
MUTED       = RGBColor(0x4A, 0x4A, 0x4A)   # muted grey
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEAD  = RGBColor(0x1C, 0x34, 0x54)   # table header fill (navy)
TABLE_ALT   = RGBColor(0xF0, 0xF4, 0xF8)   # alternating row

# ── Fonts ─────────────────────────────────────────────────────────────────────
FONT_BODY    = "Georgia"
FONT_HEADING = "Georgia"
FONT_MONO    = "Courier New"

SRC = Path(__file__).parent / "CITL_AI_Workforce_Grant_Document.md"
OUT = Path(__file__).parent / "CITL_AI_Workforce_Grant_Document.docx"


# ─────────────────────────────────────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, sides: dict) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in sides.items():
        el = OxmlElement(f"w:{side}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    # "CITL — Renton Technical College    Page N"
    run.text = "CITL — Renton Technical College    |    Page "
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    run._r.append(instrText)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def _para_space(para, before: int = 0, after: int = 6) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def _set_line_spacing(para, spacing: float = 1.15) -> None:
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:line"),    str(int(spacing * 240)))
    spacing_el.set(qn("w:lineRule"), "auto")
    pPr.append(spacing_el)


def _apply_inline_bold(para, text: str, base_size: Pt,
                        base_color: RGBColor, font_name: str, italic: bool = False) -> None:
    """
    Parse **bold** and `code` inline markers and add appropriately styled runs.
    """
    # Split on **bold** and `code`
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = base_size
            run.font.color.rgb = base_color
            if italic:
                run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = FONT_MONO
            run.font.size = Pt(base_size.pt - 1 if base_size.pt > 10 else base_size.pt)
            run.font.color.rgb = MAROON
        else:
            if part:
                run = para.add_run(part)
                run.font.name = font_name
                run.font.size = base_size
                run.font.color.rgb = base_color
                if italic:
                    run.italic = True


def _strip_md(text: str) -> str:
    """Strip **bold** and `code` markers returning plain text."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
#  Title page
# ─────────────────────────────────────────────────────────────────────────────

def build_title_page(doc: Document) -> None:
    # Top spacer
    sp = doc.add_paragraph()
    _para_space(sp, after=48)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p, before=0, after=8)
    run = p.add_run("CITL AI-Augmented IT Workforce Development Suite")
    run.font.name = FONT_HEADING
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    run.bold = True

    # Subtitle line 1
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p2, after=6)
    r2 = p2.add_run("Grant Proposal & Program Overview")
    r2.font.name = FONT_HEADING
    r2.font.size = Pt(14)
    r2.font.color.rgb = MAROON
    r2.italic = True

    # Rule — thin horizontal line via bottom border on empty para
    rule_para = doc.add_paragraph()
    rule_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(rule_para, before=10, after=10)
    pPr = rule_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1C3454")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Institution lines
    for line in [
        "Center for Instructional Technology and Learning (CITL)",
        "Renton Technical College",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_space(p, after=4)
        r = p.add_run(line)
        r.font.name = FONT_HEADING
        r.font.size = Pt(12)
        r.font.color.rgb = NAVY
        r.bold = True

    doc.add_paragraph()  # spacer

    # "Prepared for" box — shaded paragraph
    box_para = doc.add_paragraph()
    _para_space(box_para, before=18, after=18)
    box_para.paragraph_format.left_indent  = Inches(0.5)
    box_para.paragraph_format.right_indent = Inches(0.5)
    pPr = box_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EDF2F7")
    pPr.append(shd)
    lines_in_box = [
        ("Prepared for: ", "State Education and Technology Grant Programs, WA Office of Superintendent of Public Instruction (OSPI), Washington State Board for Community and Technical Colleges (SBCTC)"),
        ("Document purpose: ", "Demonstrate the scope, workforce alignment, and economic necessity of the CITL AI Workforce Training Application Suite — twelve locally-deployed, privacy-respecting, AI-powered tools that give students hands-on professional experience with skills in highest demand across Washington State's IT sector."),
    ]
    for label, body in lines_in_box:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        _para_space(p, after=6)
        pPr = p._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"),   "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"),  "EDF2F7")
        pPr.append(shd2)
        r_lbl = p.add_run(label)
        r_lbl.bold = True
        r_lbl.font.name = FONT_BODY
        r_lbl.font.size = Pt(10)
        r_lbl.font.color.rgb = NAVY
        r_body = p.add_run(body)
        r_body.font.name = FONT_BODY
        r_body.font.size = Pt(10)
        r_body.font.color.rgb = BODY_TEXT

    # Page break
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
#  Section heading styles
# ─────────────────────────────────────────────────────────────────────────────

def add_part_heading(doc: Document, text: str) -> None:
    """Part I / Part II / Part III etc — large navy, bold, ruled."""
    p = doc.add_paragraph()
    _para_space(p, before=18, after=4)
    run = p.add_run(text.upper())
    run.font.name = FONT_HEADING
    run.font.size = Pt(13)
    run.font.color.rgb = MAROON
    run.bold = True
    # Bottom border rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7B0C1A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc: Document, text: str) -> None:
    """## heading — navy, 12pt bold."""
    p = doc.add_paragraph()
    _para_space(p, before=14, after=3)
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.bold = True


def add_h3_app(doc: Document, text: str) -> None:
    """APP N: title — maroon accent box style."""
    p = doc.add_paragraph()
    _para_space(p, before=16, after=4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1C3454")
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.1)
    run = p.add_run("  " + text + "  ")
    run.font.name = FONT_HEADING
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    run.bold = True


def add_h3(doc: Document, text: str) -> None:
    """### heading — navy, italic, 11pt."""
    p = doc.add_paragraph()
    _para_space(p, before=10, after=2)
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.bold = True
    run.italic = False


# ─────────────────────────────────────────────────────────────────────────────
#  Body paragraph
# ─────────────────────────────────────────────────────────────────────────────

def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_space(p, after=6)
    _set_line_spacing(p, 1.15)
    _apply_inline_bold(p, text, Pt(11), BODY_TEXT, FONT_BODY)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _para_space(p, after=3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    _apply_inline_bold(p, text, Pt(10.5), BODY_TEXT, FONT_BODY)


def add_resume_keywords(doc: Document, text: str) -> None:
    """Styled keyword block — monospace, indented."""
    p = doc.add_paragraph()
    _para_space(p, before=4, after=8)
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F5F0E8")
    pPr.append(shd)
    label = p.add_run("Resume/Job-Board Keywords:  ")
    label.bold = True
    label.font.name = FONT_BODY
    label.font.size = Pt(9.5)
    label.font.color.rgb = NAVY
    kw_run = p.add_run(text)
    kw_run.font.name = FONT_MONO
    kw_run.font.size = Pt(9)
    kw_run.font.color.rgb = MAROON


# ─────────────────────────────────────────────────────────────────────────────
#  Table builder
# ─────────────────────────────────────────────────────────────────────────────

def add_table(doc: Document, rows_data: list[list[str]], header: bool = True) -> None:
    if not rows_data:
        return
    n_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Inches(6.5 / n_cols)

            # Background
            if r_idx == 0 and header:
                _set_cell_bg(cell, TABLE_HEAD)
            elif r_idx % 2 == 1:
                _set_cell_bg(cell, TABLE_ALT)

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _para_space(para, before=3, after=3)
            para.paragraph_format.left_indent = Inches(0.06)

            text = _strip_md(cell_text)
            run = para.add_run(text)
            run.font.name = FONT_BODY
            run.font.size = Pt(9.5)
            if r_idx == 0 and header:
                run.bold = True
                run.font.color.rgb = WHITE
            else:
                run.font.color.rgb = BODY_TEXT

    # Space after table
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
#  Markdown parser → docx writer
# ─────────────────────────────────────────────────────────────────────────────

def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Read a Markdown table starting at `start`. Returns (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        line = lines[i].strip()
        if re.fullmatch(r"[\|\-: ]+", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


APP_HEADING_RE = re.compile(
    r"### APP\s+(\d+):\s*(.*)", re.IGNORECASE
)

def is_resume_keyword_line(text: str) -> tuple[bool, str]:
    """Detect `Resume keywords this supports:` followed by backtick list."""
    if "resume keywords" in text.lower() and "`" in text:
        # Extract everything after the colon
        m = re.search(r":\s*(.+)", text)
        return True, (m.group(1) if m else text)
    return False, ""


def build_document(src: Path, out: Path) -> None:
    doc = Document()

    # ── Page layout ───────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin   = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Default Normal style ──────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_TEXT

    _add_page_number_footer(doc)
    build_title_page(doc)

    md_text = src.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # H1 — already on title page; skip
        if line.startswith("# ") and not line.startswith("## "):
            i += 1
            continue

        # HR ---
        if re.fullmatch(r"-{3,}", line.strip()):
            p = doc.add_paragraph()
            _para_space(p, before=8, after=8)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ## H2 heading
        if line.startswith("## ") and not line.startswith("### "):
            text = line[3:].strip()
            # Detect "Part I / II / III" prefix
            if re.match(r"Part\s+[IVX]+", text, re.IGNORECASE):
                add_part_heading(doc, text)
            else:
                add_h2(doc, text)
            i += 1
            continue

        # ### APP N: ... heading
        m_app = APP_HEADING_RE.match(line)
        if m_app:
            add_h3_app(doc, line[4:].strip())
            i += 1
            continue

        # ### plain H3 heading
        if line.startswith("### "):
            add_h3(doc, line[4:].strip())
            i += 1
            continue

        # Table
        if line.startswith("|"):
            table_rows, next_i = parse_table(lines, i)
            add_table(doc, table_rows)
            i = next_i
            continue

        # Bullet list
        if line.startswith("- "):
            text = line[2:].strip()
            # Detect resume keyword lines
            is_kw, kw_text = is_resume_keyword_line(text)
            if is_kw:
                # Extract just the comma-sep keywords
                kw_clean = re.sub(r"`", "", kw_text).strip()
                add_resume_keywords(doc, kw_clean)
            else:
                add_bullet(doc, text)
            i += 1
            continue

        # Blockquote — skip (already on title page)
        if line.startswith(">"):
            i += 1
            continue

        # Heading inside blockquote or metadata lines
        if line.startswith("*Document") or line.startswith("*Prepared"):
            i += 1
            continue

        # Bold-only lines (often subsection labels within body)
        if re.fullmatch(r"\*\*[^*]+\*\*:?", line.strip()):
            text = line.strip().strip("*").rstrip(":")
            p = doc.add_paragraph()
            _para_space(p, before=6, after=2)
            r = p.add_run(text + ":")
            r.font.name = FONT_HEADING
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            r.bold = True
            i += 1
            continue

        # Italic line (* text *)
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            text = line.strip().strip("*")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _para_space(p, after=4)
            r = p.add_run(text)
            r.italic = True
            r.font.name = FONT_BODY
            r.font.size = Pt(9.5)
            r.font.color.rgb = MUTED
            i += 1
            continue

        # Regular body paragraph
        text = line.strip()
        if text:
            add_body(doc, text)
        i += 1

    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_document(SRC, OUT)
