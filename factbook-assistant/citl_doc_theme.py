#!/usr/bin/env python3
"""
citl_doc_theme.py
CITL document print theme: colors, fonts, python-docx style application,
and font installation from the E READER REPO font pack.
"""
from __future__ import annotations
import ctypes, os, re, shutil, sys
from pathlib import Path
from typing import Optional

# ── Font pack location ────────────────────────────────────────────────────────
FONT_PACK = Path(r"C:\00 HENOSIS CODING PROJECTS\E READER REPO\fonts\reader-pack")

# Font family names as stored in the TTF files / used in python-docx
FONT_BODY     = "Berthold Baskerville"   # BertholdBaskerville.ttf
FONT_BODY_B   = "Berthold Baskerville"   # + bold=True
FONT_BODY_I   = "Berthold Baskerville"   # + italic=True
FONT_HEAD     = "Cheltenham"             # Cheltenham Book.ttf / Bold.ttf
FONT_CAPTION  = "Franklin Gothic Book"   # FranklinGothic Regular.ttf
FONT_MONO     = "Courier New"
FONT_FALLBACK = "Georgia"               # always present on Windows

# Map each logical name to its TTF file in the pack
FONT_FILES = {
    FONT_BODY:    [
        "BertholdBaskerville.ttf",
        "BertholdBaskerville-Bold.ttf",
        "BertholdBaskerville-Italic.ttf",
        "BertholdBaskerville-Book Italic.ttf",
    ],
    FONT_HEAD:    [
        "Cheltenham Book.ttf",
        "Cheltenham Bold.ttf",
        "Cheltenham BookItalic.ttf",
        "Cheltenham Italic.ttf",
    ],
    FONT_CAPTION: [
        "FranklinGothic Regular.ttf",
        "FranklinGothic Bold.ttf",
        "FranklinGothic Italic.ttf",
        "FranklinGothic Bold Italic.ttf",
    ],
}

# ── CITL print color palette ──────────────────────────────────────────────────
# Used as RGBColor(r, g, b) in python-docx calls
class _Palette:
    RED_ORANGE  = (0xCC, 0x33, 0x00)   # #CC3300  primary accent / rules
    SLATE_DARK  = (0x33, 0x4D, 0x6E)   # #334D6E  H1 / cover
    SLATE_MED   = (0x6B, 0x7F, 0x94)   # #6B7F94  H2 / sub-headers
    SLATE_LIGHT = (0xF0, 0xF3, 0xF6)   # #F0F3F6  callout background
    BODY_BLACK  = (0x1A, 0x1A, 0x1A)   # #1A1A1A  body text
    WHITE       = (0xFF, 0xFF, 0xFF)
    CAPTION     = (0x55, 0x55, 0x55)   # #555555  captions / footnotes
    RULE_HEX    = "CC3300"             # no-hash hex for XML attributes
    COVER_BG    = "334D6E"             # cover page header band

PAL = _Palette()

# ── Font detection (Windows registry, no pywin32 dep) ────────────────────────
def is_font_installed(family: str) -> bool:
    if sys.platform != "win32":
        return False
    import winreg
    keys = [
        (winreg.HKEY_LOCAL_MACHINE,
         r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
        (winreg.HKEY_CURRENT_USER,
         r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
    ]
    for hive, path in keys:
        try:
            with winreg.OpenKey(hive, path) as k:
                i = 0
                while True:
                    try:
                        name, _, _ = winreg.EnumValue(k, i)
                        if family.lower() in name.lower():
                            return True
                        i += 1
                    except OSError:
                        break
        except OSError:
            continue
    return False


def install_citl_fonts(log=print) -> dict:
    """
    Install all CITL reader-pack fonts for the current user (no admin needed).
    Copies TTFs to %LOCALAPPDATA%\\Microsoft\\Windows\\Fonts and registers them.
    Returns {font_file: True/False} success map.
    """
    if sys.platform != "win32":
        log("[SKIP] Font install is Windows-only.")
        return {}
    if not FONT_PACK.exists():
        log(f"[WARN] Font pack not found: {FONT_PACK}")
        return {}

    import winreg
    font_dir = Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "Windows" / "Fonts"
    font_dir.mkdir(parents=True, exist_ok=True)

    results = {}
    all_ttfs = [f for fam in FONT_FILES.values() for f in fam]
    for ttf_name in all_ttfs:
        src = FONT_PACK / ttf_name
        if not src.exists():
            log(f"[MISS] {ttf_name}")
            results[ttf_name] = False
            continue
        dst = font_dir / ttf_name
        try:
            shutil.copy2(src, dst)
            # Register in user font registry
            with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts",
                access=winreg.KEY_SET_VALUE,
            ) as k:
                label = f"{src.stem} (TrueType)"
                winreg.SetValueEx(k, label, 0, winreg.REG_SZ, str(dst))
            # Notify GDI (temporary session load)
            ctypes.windll.gdi32.AddFontResourceW(str(dst))
            ctypes.windll.user32.SendMessageW(0xFFFF, 0x001D, 0, 0)
            log(f"[OK]   {ttf_name}")
            results[ttf_name] = True
        except Exception as exc:
            log(f"[ERR]  {ttf_name}: {exc}")
            results[ttf_name] = False
    return results


def resolve_font(preferred: str) -> str:
    """Return preferred font name if installed, else FONT_FALLBACK."""
    return preferred if is_font_installed(preferred) else FONT_FALLBACK


# ── python-docx style application ────────────────────────────────────────────
def apply_citl_styles(doc) -> None:
    """
    Apply CITL print styles to a python-docx Document.
    Must be called before adding any content.
    """
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body_font    = resolve_font(FONT_BODY)
    heading_font = resolve_font(FONT_HEAD)
    caption_font = resolve_font(FONT_CAPTION)

    # ---- Page layout --------------------------------------------------------
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.18)
    sec.right_margin  = Cm(2.54)

    # ---- Normal (body) ------------------------------------------------------
    normal = doc.styles["Normal"]
    nf = normal.font
    nf.name  = body_font
    nf.size  = Pt(11)
    nf.color.rgb = RGBColor(*PAL.BODY_BLACK)
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # ---- Heading 1 ----------------------------------------------------------
    h1 = doc.styles["Heading 1"]
    h1f = h1.font
    h1f.name  = heading_font
    h1f.bold  = True
    h1f.size  = Pt(18)
    h1f.color.rgb = RGBColor(*PAL.SLATE_DARK)
    h1f.underline = False
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # ---- Heading 2 ----------------------------------------------------------
    h2 = doc.styles["Heading 2"]
    h2f = h2.font
    h2f.name  = heading_font
    h2f.bold  = False
    h2f.size  = Pt(14)
    h2f.color.rgb = RGBColor(*PAL.SLATE_MED)
    h2f.underline = False
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ---- Heading 3 ----------------------------------------------------------
    h3 = doc.styles["Heading 3"]
    h3f = h3.font
    h3f.name   = caption_font
    h3f.bold   = True
    h3f.italic = False
    h3f.size   = Pt(11)
    h3f.color.rgb = RGBColor(*PAL.BODY_BLACK)
    h3f.underline = False
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # ---- Caption style ------------------------------------------------------
    try:
        cap = doc.styles["Caption"]
    except KeyError:
        cap = doc.styles.add_style("Caption", 1)
    cap.font.name  = caption_font
    cap.font.size  = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(*PAL.CAPTION)
    cap.paragraph_format.space_after = Pt(8)

    # ---- Header / Footer ----------------------------------------------------
    _build_header(doc, heading_font, caption_font)
    _build_footer(doc, caption_font)


def _build_header(doc, heading_font: str, caption_font: str) -> None:
    from docx.shared import Pt, RGBColor, Tab
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header = doc.sections[0].header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    if not header.paragraphs:
        header.add_paragraph()

    para = header.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_citl = para.add_run("CITL  ")
    run_citl.font.name  = heading_font
    run_citl.font.bold  = True
    run_citl.font.size  = Pt(9)
    run_citl.font.color.rgb = RGBColor(*PAL.RED_ORANGE)

    run_title = para.add_run("Center for Information Technology and Learning")
    run_title.font.name  = caption_font
    run_title.font.size  = Pt(9)
    run_title.font.color.rgb = RGBColor(*PAL.CAPTION)

    # Red-orange bottom rule on the header paragraph
    _add_para_border(para, side="bottom", color=PAL.RULE_HEX, sz="6")


def _build_footer(doc, caption_font: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        footer.add_paragraph()

    para = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_border(para, side="top", color=PAL.RULE_HEX, sz="4")

    # "Page N" field
    run = para.add_run()
    run.font.name  = caption_font
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(*PAL.CAPTION)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = para.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

    para.add_run("  ·  CITL Documentation").font.size = Pt(9)


# ── DOCX building helpers ─────────────────────────────────────────────────────
def add_rule(doc, color_hex: str = PAL.RULE_HEX) -> None:
    """Add a thin colored horizontal rule paragraph."""
    from docx.shared import Pt
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    _add_para_border(para, side="bottom", color=color_hex, sz="6")


def add_h1_with_bar(doc, text: str) -> None:
    """H1 heading with a left red-orange bar."""
    para = doc.add_heading(text, level=1)
    _add_para_border(para, side="left",
                     color=PAL.RULE_HEX, sz="20", space="6")


def add_h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc, text: str) -> None:
    doc.add_heading(text, level=3)


def add_body(doc, text: str) -> None:
    """Add a Normal-style paragraph, auto-detecting numbered steps."""
    from docx.shared import Pt, Inches
    text = text.strip()
    if not text:
        return
    # Detect numbered step: "1." or "Step 1:"
    if re.match(r"^(\d+\.|\bStep\s+\d+:)", text):
        p = doc.add_paragraph(style="List Number")
        p.text = re.sub(r"^(\d+\.|\bStep\s+\d+:)\s*", "", text)
    elif re.match(r"^[-•]\s", text):
        p = doc.add_paragraph(style="List Bullet")
        p.text = text[2:]
    else:
        doc.add_paragraph(text)


def add_callout(doc, text: str, kind: str = "note") -> None:
    """
    Add a shaded callout box.
    kind: "note" | "warning" | "tip"
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    labels = {"note": "NOTE", "warning": "WARNING", "tip": "TIP"}
    label_colors = {
        "note":    PAL.SLATE_MED,
        "warning": (0xCC, 0x44, 0x00),
        "tip":     (0x22, 0x77, 0x44),
    }
    label = labels.get(kind, "NOTE")
    lcolor = label_colors.get(kind, PAL.SLATE_MED)

    # Use a 1-column table for the shaded box
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "F0F3F6")
    _set_cell_border(cell, PAL.RULE_HEX if kind == "warning" else "6B7F94")

    # Label run
    p = cell.paragraphs[0]
    p.clear()
    label_run = p.add_run(label + "  ")
    label_run.font.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(*lcolor)
    label_run.font.name = resolve_font(FONT_CAPTION)

    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.name = resolve_font(FONT_BODY)
    body_run.font.color.rgb = RGBColor(*PAL.BODY_BLACK)
    doc.add_paragraph()   # spacing after


def add_screenshot_placeholder(doc, caption: str = "", lines: int = 7) -> None:
    """
    Add a bordered screenshot placeholder box so editors can paste captures later.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "FBFCFD")
    _set_cell_border(cell, "6B7F94")

    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head = p.add_run("SCREENSHOT PLACEHOLDER")
    head.font.bold = True
    head.font.size = Pt(10)
    head.font.name = resolve_font(FONT_CAPTION)
    head.font.color.rgb = RGBColor(*PAL.SLATE_MED)

    cap = caption.strip() or "Paste UI capture here."
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = p2.add_run(cap)
    c_run.font.size = Pt(9)
    c_run.font.name = resolve_font(FONT_CAPTION)
    c_run.font.color.rgb = RGBColor(*PAL.CAPTION)

    for _ in range(max(4, int(lines))):
        cell.add_paragraph(" ")

    doc.add_paragraph()


def add_cover_page(doc, meta: dict) -> None:
    """
    Generate a styled CITL cover page.
    meta keys: title, subtitle, app_name, version, author, date, template_type
    """
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    heading_font = resolve_font(FONT_HEAD)
    body_font    = resolve_font(FONT_BODY)
    caption_font = resolve_font(FONT_CAPTION)

    def _centered(text, font, size, bold=False, color=PAL.BODY_BLACK, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name  = font
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.color.rgb = RGBColor(*color)
        return p

    # Top spacer
    for _ in range(4):
        doc.add_paragraph()

    # CITL label
    _centered("CENTER FOR INFORMATION TECHNOLOGY AND LEARNING",
              caption_font, 9, color=PAL.RED_ORANGE, space_after=2)

    # Red rule
    add_rule(doc)
    doc.add_paragraph()

    # Document type
    _centered(meta.get("template_type", "Technical Document").upper(),
              caption_font, 10, color=PAL.SLATE_MED, space_after=18)

    # App name
    _centered(meta.get("app_name", "CITL Application"),
              heading_font, 28, bold=True, color=PAL.SLATE_DARK, space_after=6)

    # Title
    _centered(meta.get("title", ""),
              heading_font, 18, color=PAL.RED_ORANGE, space_after=4)

    # Subtitle
    if meta.get("subtitle"):
        _centered(meta["subtitle"], body_font, 12,
                  color=PAL.SLATE_MED, space_after=0)

    doc.add_paragraph()
    add_rule(doc)
    doc.add_paragraph()

    # Metadata block
    for label, key in [("Version", "version"), ("Author", "author"), ("Date", "date")]:
        val = meta.get(key, "")
        if val:
            _centered(f"{label}:  {val}", caption_font, 10, color=PAL.CAPTION)

    # Page break
    doc.add_page_break()


# ── Internal XML helpers ──────────────────────────────────────────────────────
def _add_para_border(para, side: str, color: str,
                     sz: str = "6", space: str = "1") -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr  = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    sz)
    bdr.set(qn("w:space"), space)
    bdr.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bdr)
    pPr.append(pBdr)


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border(cell, color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ("top", "bottom", "left", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"),   "single")
        bdr.set(qn("w:sz"),    "4")
        bdr.set(qn("w:space"), "0")
        bdr.set(qn("w:color"), color.lstrip("#"))
        tcBdr.append(bdr)
    tcPr.append(tcBdr)
