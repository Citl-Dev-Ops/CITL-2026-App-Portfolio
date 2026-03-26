#!/usr/bin/env python3
"""
CITL Document Composer v1.1
Professional technical manual and walkthrough generator for CITL apps.
"""
from __future__ import annotations

import os
import queue
import re
import subprocess
import sys
import threading
import traceback
from datetime import date
from pathlib import Path
from typing import List, Optional

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext, ttk
except ImportError:
    print("tkinter required.")
    sys.exit(1)

from citl_doc_theme import (
    FONT_BODY,
    FONT_CAPTION,
    FONT_HEAD,
    add_body,
    add_callout,
    add_cover_page,
    add_h1_with_bar,
    add_h3,
    add_rule,
    add_screenshot_placeholder,
    apply_citl_styles,
    install_citl_fonts,
    is_font_installed,
)
from citl_doc_templates import (
    TEMPLATE_NAMES,
    get_best_model,
    get_best_vision_model,
    get_ollama_models,
    get_sections,
    stream_generate,
)

_HERE = Path(__file__).parent
if getattr(sys, "frozen", False):
    _env_repo = os.environ.get("CITL_REPO", "").strip()
    if _env_repo and Path(_env_repo).is_dir():
        REPO = Path(_env_repo)
    else:
        REPO = Path(sys.executable).parent.parent.parent
else:
    REPO = _HERE.parent
DOCS_DIR = REPO / "documents"

C = {
    "bg": "#140a0a",
    "panel": "#1e0f0f",
    "panel_alt": "#271414",
    "card_sel": "#521c1c",
    "text": "#f5eeee",
    "muted": "#c4a0a0",
    "faint": "#8a7070",
    "accent": "#d84444",
    "btn": "#4a1a1a",
    "btn_hi": "#6e2525",
    "btn_acc": "#7a1e1e",
    "notebk": "#180c0c",
}
_F = "Segoe UI" if sys.platform == "win32" else "Ubuntu"

APP_NAME = "CITL Document Composer"
APP_VERSION = "v1.1"

_RE_NUM = re.compile(r"^(\d+(?:\.\d+){0,4})[\).:]?\s+(.+)$")
_RE_ALPHA = re.compile(r"^([a-zA-Z])[\).:]\s+(.+)$")
_RE_BUL = re.compile(r"^(\s*)[-*\u2022]\s+(.+)$")
_RE_CALL = re.compile(r"^(TIP|NOTE|WARNING):\s*(.+)$", re.IGNORECASE)
_RE_SHOT = re.compile(r"^(SCREENSHOT(?:\s+PLACEHOLDER)?|IMAGE):\s*(.+)$", re.IGNORECASE)
_RE_SUBHD = re.compile(
    r"^(Menu Path|Expected Result|Validation|Troubleshooting|Context)\s*:\s*(.*)$",
    re.IGNORECASE,
)


def _list_style_name(ordered: bool, level: int) -> str:
    idx = max(1, min(3, level + 1))
    if ordered:
        return "List Number" if idx == 1 else f"List Number {idx}"
    return "List Bullet" if idx == 1 else f"List Bullet {idx}"


def _add_list_line(doc, text: str, ordered: bool, level: int) -> None:
    text = text.strip()
    if not text:
        return
    style = _list_style_name(ordered, level)
    try:
        p = doc.add_paragraph(style=style)
        p.text = text
    except Exception:
        prefix = f"{level + 1}. " if ordered else "- "
        doc.add_paragraph(prefix + text)


def _render_structured_content(doc, content: str) -> None:
    para_buf: List[str] = []

    def _flush_paragraph() -> None:
        if para_buf:
            add_body(doc, " ".join(para_buf).strip())
            para_buf.clear()

    for raw in content.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            _flush_paragraph()
            continue

        m = _RE_CALL.match(stripped)
        if m:
            _flush_paragraph()
            add_callout(doc, m.group(2).strip(), m.group(1).lower())
            continue

        m = _RE_SHOT.match(stripped)
        if m:
            _flush_paragraph()
            add_screenshot_placeholder(doc, m.group(2).strip() or "Capture this step.")
            continue

        m = _RE_SUBHD.match(stripped)
        if m:
            _flush_paragraph()
            add_h3(doc, m.group(1).strip().title())
            if m.group(2).strip():
                add_body(doc, m.group(2).strip())
            continue

        m = _RE_NUM.match(stripped)
        if m:
            _flush_paragraph()
            depth = min(2, m.group(1).count("."))
            _add_list_line(doc, m.group(2).strip(), ordered=True, level=depth)
            continue

        m = _RE_ALPHA.match(stripped)
        if m:
            _flush_paragraph()
            _add_list_line(doc, m.group(2).strip(), ordered=True, level=1)
            continue

        m = _RE_BUL.match(line)
        if m:
            _flush_paragraph()
            indent = len(m.group(1).replace("\t", "    "))
            depth = min(2, indent // 2)
            _add_list_line(doc, m.group(2).strip(), ordered=False, level=depth)
            continue

        para_buf.append(stripped)

    _flush_paragraph()


def _export_docx(sections: List[dict], meta: dict, out_path: str) -> None:
    from docx import Document

    doc = Document()
    apply_citl_styles(doc)

    for sec in sections:
        if sec["id"] == "cover":
            add_cover_page(doc, meta)
            continue

        add_h1_with_bar(doc, sec["title"])
        content = sec.get("content", "").strip()
        if not content:
            add_body(doc, f"[{sec['title']} - content not yet generated]")
            add_rule(doc)
            continue

        _render_structured_content(doc, content)
        add_rule(doc)

    doc.save(out_path)


class DocComposer:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title(f"{APP_NAME}  {APP_VERSION}")
        self.root.configure(bg=C["bg"])
        self.root.minsize(1060, 700)

        self._sections: List[dict] = []
        self._current_idx = -1
        self._generating = False
        self._gen_stop = threading.Event()

        self._models: List[dict] = []
        self._model_var = tk.StringVar()
        self._model_disp = tk.StringVar(value="Checking Ollama...")

        self._template_var = tk.StringVar(value=TEMPLATE_NAMES[0])
        self._app_name_var = tk.StringVar(value="CITL Factbook Assistant")
        self._version_var = tk.StringVar(value="v1.0")
        self._author_var = tk.StringVar(value="Abdo Mohammad / CITL")
        self._date_var = tk.StringVar(value=str(date.today()))
        self._subtitle_var = tk.StringVar()
        self._ui_goal_var = tk.StringVar()

        self._status_var = tk.StringVar(value="Ready")
        self._font_status = tk.StringVar(value="Checking fonts...")
        self._shot_count_var = tk.StringVar(value="Screenshots: 0 attached")
        self._screenshot_paths: List[Path] = []
        self._warned_non_vision = False

        DOCS_DIR.mkdir(parents=True, exist_ok=True)

        self._build_ui()
        self._on_template_change()
        self.root.after(350, self._detect_models)
        self.root.after(500, self._check_fonts)

    def _build_ui(self) -> None:
        hdr = tk.Frame(self.root, bg=C["panel"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["accent"], height=4).pack(fill="x")
        hi = tk.Frame(hdr, bg=C["panel"])
        hi.pack(fill="x", padx=18, pady=10)
        tk.Label(hi, text=APP_NAME, font=(_F, 18, "bold"),
                 bg=C["panel"], fg=C["text"]).pack(side="left")
        tk.Label(hi, text=APP_VERSION, font=(_F, 10, "bold"),
                 bg=C["panel"], fg=C["accent"]).pack(side="left", padx=6)
        tk.Label(hi, textvariable=self._model_disp, font=(_F, 9, "italic"),
                 bg=C["panel"], fg=C["muted"]).pack(side="right")

        sb = tk.Frame(self.root, bg=C["panel_alt"])
        sb.pack(fill="x")
        tk.Label(sb, textvariable=self._status_var, font=(_F, 8),
                 bg=C["panel_alt"], fg=C["accent"]).pack(side="left", padx=10)
        tk.Label(sb, textvariable=self._font_status, font=(_F, 8),
                 bg=C["panel_alt"], fg=C["muted"]).pack(side="right", padx=10)

        body = tk.Frame(self.root, bg=C["bg"])
        body.pack(fill="both", expand=True, padx=8, pady=6)
        body.columnconfigure(0, weight=0, minsize=210)
        body.columnconfigure(1, weight=1)
        body.columnconfigure(2, weight=0, minsize=280)
        body.rowconfigure(0, weight=1)

        self._build_left(body)
        self._build_center(body)
        self._build_right(body)

    def _build_left(self, parent) -> None:
        left = tk.Frame(parent, bg=C["panel"], relief="flat", bd=0)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 6))

        self._lbl(left, "TEMPLATE").pack(anchor="w", padx=8, pady=(8, 2))
        self._tmpl_combo = ttk.Combobox(
            left,
            textvariable=self._template_var,
            values=TEMPLATE_NAMES,
            state="readonly",
            font=(_F, 9),
            width=26,
        )
        self._tmpl_combo.pack(fill="x", padx=8, pady=(0, 8))
        self._tmpl_combo.bind("<<ComboboxSelected>>", self._on_template_change)

        self._lbl(left, "SECTIONS").pack(anchor="w", padx=8, pady=(4, 2))
        self._section_lb = tk.Listbox(
            left,
            font=(_F, 9),
            bg=C["notebk"],
            fg=C["text"],
            selectbackground=C["card_sel"],
            selectforeground=C["accent"],
            activestyle="none",
            bd=0,
            highlightthickness=0,
            relief="flat",
            selectmode="single",
        )
        self._section_lb.pack(fill="both", expand=True, padx=8, pady=(0, 4))
        self._section_lb.bind("<<ListboxSelect>>", self._on_section_select)

        self._gen_all_btn = self._btn(left, "Generate All Sections", C["btn_acc"], self._generate_all)
        self._gen_all_btn.pack(fill="x", padx=8, pady=(0, 4))

        self._stop_btn = self._btn(left, "Stop", C["btn"], self._stop_generate)
        self._stop_btn.pack(fill="x", padx=8, pady=(0, 8))
        self._stop_btn.config(state="disabled")

    def _build_center(self, parent) -> None:
        center = tk.Frame(parent, bg=C["bg"])
        center.grid(row=0, column=1, sticky="nsew", padx=6)
        center.rowconfigure(1, weight=1)
        center.columnconfigure(0, weight=1)

        self._sec_label = tk.Label(
            center,
            text="Select a section",
            font=(_F, 12, "bold"),
            bg=C["bg"],
            fg=C["accent"],
            anchor="w",
        )
        self._sec_label.grid(row=0, column=0, sticky="ew", pady=(0, 4))

        self._editor = scrolledtext.ScrolledText(
            center,
            font=("Georgia", 10),
            bg=C["notebk"],
            fg="#e8e0d4",
            insertbackground=C["text"],
            relief="flat",
            wrap="word",
            state="disabled",
        )
        self._editor.grid(row=1, column=0, sticky="nsew")

        ec = tk.Frame(center, bg=C["bg"])
        ec.grid(row=2, column=0, sticky="ew", pady=(4, 0))
        self._btn(ec, "Generate This Section", C["btn_acc"], self._generate_section).pack(side="left", padx=(0, 6))
        self._btn(ec, "Clear Section", C["btn"], self._clear_section).pack(side="left")
        tk.Label(
            ec,
            text="(edit freely; content is saved as you type)",
            font=(_F, 8, "italic"),
            bg=C["bg"],
            fg=C["faint"],
        ).pack(side="right")

        self._editor.bind("<KeyRelease>", self._on_editor_change)

    def _build_right(self, parent) -> None:
        right = tk.Frame(parent, bg=C["panel"])
        right.grid(row=0, column=2, sticky="nsew")
        right.columnconfigure(1, weight=1)

        def _row(label, var, row):
            tk.Label(right, text=label, font=(_F, 8), bg=C["panel"], fg=C["muted"], anchor="w").grid(
                row=row, column=0, sticky="ew", padx=(8, 4), pady=2
            )
            tk.Entry(
                right,
                textvariable=var,
                font=(_F, 9),
                bg=C["notebk"],
                fg=C["text"],
                insertbackground=C["text"],
                relief="flat",
            ).grid(row=row, column=1, sticky="ew", padx=(0, 8), pady=2)

        row = 0
        self._lbl(right, "DOCUMENT INFO").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(8, 4)); row += 1
        _row("App name:", self._app_name_var, row); row += 1
        _row("Version:", self._version_var, row); row += 1
        _row("Author:", self._author_var, row); row += 1
        _row("Date:", self._date_var, row); row += 1
        _row("Subtitle:", self._subtitle_var, row); row += 1

        self._lbl(right, "TOPIC / PROMPT").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(10, 2)); row += 1
        tk.Label(
            right,
            text="Describe what this document covers. The AI uses this for every section.",
            font=(_F, 8, "italic"),
            bg=C["panel"],
            fg=C["faint"],
            justify="left",
            wraplength=240,
        ).grid(row=row, column=0, columnspan=2, sticky="ew", padx=8); row += 1
        self._topic_text = tk.Text(
            right, font=(_F, 9), bg=C["notebk"], fg=C["text"],
            insertbackground=C["text"], relief="flat", height=4, wrap="word"
        )
        self._topic_text.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(2, 8)); row += 1

        self._lbl(right, "MENU DIVING TARGET").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2)); row += 1
        _row("UI goal:", self._ui_goal_var, row); row += 1

        self._lbl(right, "SCREENSHOT EVIDENCE").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2)); row += 1
        tk.Label(right, textvariable=self._shot_count_var, font=(_F, 8), bg=C["panel"], fg=C["muted"], justify="left").grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8
        ); row += 1
        shot_btns = tk.Frame(right, bg=C["panel"])
        shot_btns.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(2, 4))
        shot_btns.columnconfigure(0, weight=1)
        shot_btns.columnconfigure(1, weight=1)
        self._btn(shot_btns, "Attach screenshots", C["btn"], self._add_screenshots).grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self._btn(shot_btns, "Clear screenshots", C["btn"], self._clear_screenshots).grid(row=0, column=1, sticky="ew", padx=(4, 0))
        row += 1
        self._shot_notes_text = tk.Text(
            right, font=(_F, 8), bg=C["notebk"], fg=C["text"],
            insertbackground=C["text"], relief="flat", height=4, wrap="word"
        )
        self._shot_notes_text.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 8)); row += 1
        self._shot_notes_text.insert(
            "1.0",
            "Optional: describe labels/buttons visible in screenshots.",
        )

        self._lbl(right, "OLLAMA MODEL").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2)); row += 1
        self._model_combo = ttk.Combobox(right, textvariable=self._model_var, font=(_F, 9), state="readonly", width=28)
        self._model_combo.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 2)); row += 1
        self._btn(right, "Refresh Models", C["btn"], self._detect_models).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 8)
        ); row += 1

        self._lbl(right, "TYPOGRAPHY").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2)); row += 1
        tk.Label(
            right, textvariable=self._font_status, font=(_F, 8),
            bg=C["panel"], fg=C["muted"], wraplength=240, justify="left"
        ).grid(row=row, column=0, columnspan=2, sticky="ew", padx=8); row += 1
        self._btn(right, "Install CITL Fonts", C["btn"], self._install_fonts).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 8)
        ); row += 1

        self._lbl(right, "EXPORT").grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2)); row += 1
        self._export_btn = self._btn(right, "Export to .docx", C["btn_acc"], self._export)
        self._export_btn.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 4)); row += 1
        self._btn(right, "Open Documents Folder", C["btn"], self._open_docs_dir).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 8)
        )

    def _lbl(self, parent, text: str) -> tk.Label:
        return tk.Label(parent, text=text, font=(_F, 8, "bold"), bg=C["panel"], fg=C["accent"])

    def _btn(self, parent, text: str, bg: str, cmd) -> tk.Button:
        return tk.Button(
            parent,
            text=text,
            font=(_F, 9),
            bg=bg,
            fg=C["text"],
            activebackground=C["btn_hi"],
            relief="flat",
            bd=0,
            padx=8,
            pady=5,
            command=cmd,
        )

    def _update_shot_status(self) -> None:
        self._shot_count_var.set(f"Screenshots: {len(self._screenshot_paths)} attached")

    def _add_screenshots(self) -> None:
        picked = filedialog.askopenfilenames(
            title="Attach screenshots",
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.bmp *.webp"),
                ("All files", "*.*"),
            ],
            initialdir=str(DOCS_DIR),
        )
        if not picked:
            return
        seen = {str(p).lower() for p in self._screenshot_paths}
        for raw in picked:
            p = Path(raw)
            k = str(p).lower()
            if k not in seen and p.exists():
                seen.add(k)
                self._screenshot_paths.append(p)
        self._update_shot_status()

    def _clear_screenshots(self) -> None:
        self._screenshot_paths = []
        self._update_shot_status()

    def _get_meta(self) -> dict:
        topic = self._topic_text.get("1.0", "end").strip()
        shot_notes = self._shot_notes_text.get("1.0", "end").strip()
        if shot_notes.lower().startswith("optional:"):
            shot_notes = ""
        return {
            "app_name": self._app_name_var.get(),
            "version": self._version_var.get(),
            "author": self._author_var.get(),
            "date": self._date_var.get(),
            "subtitle": self._subtitle_var.get(),
            "topic": topic,
            "ui_goal": self._ui_goal_var.get(),
            "screenshot_notes": shot_notes,
            "screenshot_count": str(len(self._screenshot_paths)),
            "template_type": self._template_var.get(),
            "title": self._app_name_var.get(),
        }

    def _on_template_change(self, *_):
        self._save_current_section()
        self._sections = get_sections(self._template_var.get())
        self._current_idx = -1
        self._section_lb.delete(0, "end")
        for s in self._sections:
            self._section_lb.insert("end", s["title"])
        if self._sections:
            self._section_lb.selection_set(0)
            self._load_section(0)

    def _on_section_select(self, *_):
        sel = self._section_lb.curselection()
        if not sel:
            return
        idx = int(sel[0])
        self._save_current_section()
        self._load_section(idx)

    def _load_section(self, idx: int):
        self._current_idx = idx
        sec = self._sections[idx]
        self._sec_label.config(text=sec["title"])
        self._editor.config(state="normal")
        self._editor.delete("1.0", "end")
        if sec["id"] == "cover":
            self._editor.insert("1.0", "(Cover page is generated automatically from the fields on the right.)")
            self._editor.config(state="disabled")
        else:
            self._editor.insert("1.0", sec.get("content", ""))
            self._editor.config(state="normal")

    def _save_current_section(self):
        if 0 <= self._current_idx < len(self._sections):
            sec = self._sections[self._current_idx]
            if sec["id"] != "cover":
                sec["content"] = self._editor.get("1.0", "end-1c")

    def _on_editor_change(self, *_):
        self._save_current_section()

    def _clear_section(self):
        if 0 <= self._current_idx < len(self._sections):
            self._sections[self._current_idx]["content"] = ""
            self._editor.config(state="normal")
            self._editor.delete("1.0", "end")

    def _detect_models(self):
        def _run():
            models = get_ollama_models()

            def _update():
                self._models = models
                if models:
                    names = [m["display"] for m in models]
                    self._model_combo["values"] = names
                    self._model_combo.current(0)
                    self._model_var.set(names[0])
                    best = models[0]
                    vtag = " [vision]" if best.get("is_vision") else ""
                    self._model_disp.set(f"Best model: {best['name']}{vtag} ({best['params']}B)")
                else:
                    self._model_disp.set("No Ollama models found.")

            self.root.after(0, _update)

        threading.Thread(target=_run, daemon=True).start()

    def _selected_model_entry(self) -> Optional[dict]:
        idx = self._model_combo.current()
        if 0 <= idx < len(self._models):
            return self._models[idx]
        pick = self._model_var.get().strip()
        for m in self._models:
            if pick in (m.get("display", ""), m.get("name", "")):
                return m
        return self._models[0] if self._models else None

    def _pick_generation_model(self) -> tuple[Optional[str], bool]:
        chosen = self._selected_model_entry()
        model_name = chosen["name"] if chosen else get_best_model()
        use_images = bool(self._screenshot_paths)
        if not model_name:
            return None, False

        if use_images and not (chosen and chosen.get("is_vision")):
            vision = get_best_vision_model()
            if vision:
                model_name = vision
                self._status_var.set(f"Using vision model for screenshots: {vision}")
            else:
                if not self._warned_non_vision:
                    messagebox.showwarning(
                        APP_NAME,
                        "No vision-capable Ollama model found.\n"
                        "Screenshots will be ignored for image interpretation.\n"
                        "Text notes are still used.",
                    )
                    self._warned_non_vision = True
                use_images = False
        return model_name, use_images

    def _generate_section(self, idx: Optional[int] = None, done_cb=None):
        if self._generating:
            return
        target = idx if idx is not None else self._current_idx
        if not (0 <= target < len(self._sections)):
            return
        sec = self._sections[target]
        if sec["id"] == "cover":
            if done_cb:
                done_cb()
            return

        model_name, use_images = self._pick_generation_model()
        if not model_name:
            messagebox.showwarning(APP_NAME, "No Ollama model found. Make sure Ollama is running.")
            return

        if target != self._current_idx:
            self._save_current_section()
            self._load_section(target)
            self._section_lb.selection_clear(0, "end")
            self._section_lb.selection_set(target)

        self._editor.config(state="normal")
        self._editor.delete("1.0", "end")
        self._sec_label.config(text=f"Generating: {sec['title']}...")
        self._status_var.set(f"Generating '{sec['title']}' with {model_name}...")
        self._set_generating(True)
        self._gen_stop.clear()

        q = stream_generate(
            model=model_name,
            section_prompt=sec["prompt"],
            meta=self._get_meta(),
            token_cb=lambda _t: None,
            done_cb=lambda _ok, _msg: None,
            image_paths=[str(p) for p in self._screenshot_paths] if use_images else None,
        )
        self._poll_gen(q, target, done_cb)

    def _poll_gen(self, q: queue.Queue, target: int, done_cb=None):
        if self._gen_stop.is_set():
            self._finish_gen(target, done_cb)
            return
        try:
            while True:
                msg_type, content = q.get_nowait()
                if msg_type == "token":
                    self._editor.insert("end", content)
                    self._editor.see("end")
                elif msg_type == "done":
                    self._finish_gen(target, done_cb)
                    return
                elif msg_type == "error":
                    self._status_var.set(f"Error: {content}")
                    self._finish_gen(target, done_cb)
                    return
        except queue.Empty:
            pass
        self.root.after(40, lambda: self._poll_gen(q, target, done_cb))

    def _finish_gen(self, target: int, done_cb=None):
        self._save_current_section()
        sec = self._sections[target]
        self._sec_label.config(text=sec["title"])
        self._status_var.set(f"Done: {sec['title']}")
        self._set_generating(False)
        if done_cb:
            done_cb()

    def _generate_all(self):
        if self._generating:
            return
        idxs = [i for i, s in enumerate(self._sections) if s["id"] != "cover" and not s.get("content", "").strip()]
        if not idxs:
            idxs = [i for i, s in enumerate(self._sections) if s["id"] != "cover"]
        self._gen_stop.clear()
        self._run_gen_sequence(idxs, 0)

    def _run_gen_sequence(self, idxs: List[int], pos: int):
        if pos >= len(idxs) or self._gen_stop.is_set():
            self._status_var.set("All sections generated." if pos >= len(idxs) else "Generation stopped.")
            self._set_generating(False)
            return
        idx = idxs[pos]
        self._generate_section(
            idx=idx,
            done_cb=lambda: self.root.after(200, lambda: self._run_gen_sequence(idxs, pos + 1)),
        )

    def _stop_generate(self):
        self._gen_stop.set()
        self._status_var.set("Stopping...")

    def _set_generating(self, val: bool):
        self._generating = val
        state = "disabled" if val else "normal"
        self._gen_all_btn.config(state=state)
        self._stop_btn.config(state="normal" if val else "disabled")
        self._export_btn.config(state=state)

    def _check_fonts(self):
        families = [FONT_BODY, FONT_HEAD, FONT_CAPTION]
        missing = [f for f in families if not is_font_installed(f)]
        if not missing:
            self._font_status.set(f"Installed: {FONT_BODY}, {FONT_HEAD}, {FONT_CAPTION}")
        else:
            self._font_status.set(f"Missing: {', '.join(missing)}")

    def _install_fonts(self):
        def _run():
            results = install_citl_fonts(log=lambda _msg: None)
            ok = sum(1 for v in results.values() if v)
            fail = sum(1 for v in results.values() if not v)
            self.root.after(0, lambda: self._font_status.set(
                f"Installed {ok} font(s)." + (f" {fail} failed." if fail else "")
            ))
            self.root.after(0, self._check_fonts)

        threading.Thread(target=_run, daemon=True).start()

    def _export(self):
        self._save_current_section()
        meta = self._get_meta()
        app_safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in meta["app_name"])
        default_name = f"CITL_{app_safe}_{self._template_var.get().replace(' ', '_')}.docx"
        out_path = filedialog.asksaveasfilename(
            title="Save CITL Document",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialdir=str(DOCS_DIR),
            initialfile=default_name,
        )
        if not out_path:
            return

        self._status_var.set("Exporting...")
        self.root.update_idletasks()
        try:
            _export_docx(self._sections, meta, out_path)
            self._status_var.set(f"Saved: {Path(out_path).name}")
            if messagebox.askyesno(APP_NAME, f"Document saved:\n{out_path}\n\nOpen it now?"):
                self._open_path(Path(out_path))
        except Exception as exc:
            self._status_var.set(f"Export failed: {exc}")
            messagebox.showerror(APP_NAME, f"Export failed:\n{exc}\n\nMake sure python-docx is installed:\n  pip install python-docx")

    def _open_path(self, path: Path):
        try:
            if sys.platform == "win32":
                os.startfile(str(path))
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(path)])
            else:
                subprocess.Popen(["xdg-open", str(path)])
        except Exception:
            pass

    def _open_docs_dir(self):
        DOCS_DIR.mkdir(parents=True, exist_ok=True)
        self._open_path(DOCS_DIR)


def main():
    root = tk.Tk()
    root.withdraw()
    try:
        root.tk.call("tk", "scaling", 1.25)
    except Exception:
        pass
    root.deiconify()
    try:
        DocComposer(root)
        root.mainloop()
    except Exception as exc:
        log_path = _HERE / "citl_doc_composer_crash.log"
        log_path.write_text(f"{traceback.format_exc()}\n", encoding="utf-8")
        try:
            messagebox.showerror(APP_NAME, f"{exc}\nSee: {log_path}")
        except Exception:
            print(traceback.format_exc(), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

